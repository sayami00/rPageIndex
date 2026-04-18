"""DOCX parser tests — builds documents in-memory with python-docx."""
from __future__ import annotations

import pytest

pytest.importorskip("docx")

import docx as python_docx
from docx import Document

from src.ingestion.docx_parser import DOCXParser


def _make_doc_with_headings(tmp_path) -> str:
    path = str(tmp_path / "headings.docx")
    doc = Document()
    doc.add_heading("Chapter One", level=1)
    doc.add_paragraph("Content of chapter one.")
    doc.add_heading("Chapter Two", level=1)
    doc.add_paragraph("Content of chapter two.")
    doc.add_heading("Chapter Three", level=1)
    doc.add_paragraph("Content of chapter three.")
    doc.save(path)
    return path


def _make_doc_no_headings(tmp_path) -> str:
    path = str(tmp_path / "noheadings.docx")
    doc = Document()
    doc.add_paragraph("First paragraph.")
    doc.add_paragraph("Second paragraph.")
    doc.add_paragraph("Third paragraph.")
    doc.save(path)
    return path


def _make_doc_with_table(tmp_path) -> str:
    path = str(tmp_path / "table.docx")
    doc = Document()
    doc.add_heading("Section", level=1)
    table = doc.add_table(rows=3, cols=3)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Age"
    table.cell(0, 2).text = "City"
    table.cell(1, 0).text = "Alice"
    table.cell(1, 1).text = "30"
    table.cell(1, 2).text = "NY"
    table.cell(2, 0).text = "Bob"
    table.cell(2, 1).text = "25"
    table.cell(2, 2).text = "LA"
    doc.save(path)
    return path


def _make_doc_with_list(tmp_path) -> str:
    path = str(tmp_path / "list.docx")
    doc = Document()
    doc.add_paragraph("Item one", style="List Bullet")
    doc.add_paragraph("Item two", style="List Bullet")
    doc.add_paragraph("Item three", style="List Bullet")
    doc.save(path)
    return path


def test_docx_logical_pages(tmp_path):
    path = _make_doc_with_headings(tmp_path)
    blocks = DOCXParser().parse(path)
    page_nums = sorted(set(b.page_number for b in blocks))
    assert len(page_nums) == 3, f"Expected 3 logical pages, got {page_nums}"
    assert page_nums == [1, 2, 3]


def test_docx_no_headings_single_page(tmp_path):
    path = _make_doc_no_headings(tmp_path)
    blocks = DOCXParser().parse(path)
    assert len(blocks) > 0
    assert all(b.page_number == 1 for b in blocks)


def test_docx_table_extraction(tmp_path):
    path = _make_doc_with_table(tmp_path)
    blocks = DOCXParser().parse(path)
    table_blocks = [b for b in blocks if b.block_type_hint == "table"]
    assert len(table_blocks) >= 1
    tb = table_blocks[0]
    assert tb.raw_headers is not None
    assert "Name" in tb.raw_headers
    assert len(tb.raw_rows) >= 1


def test_docx_list_items(tmp_path):
    path = _make_doc_with_list(tmp_path)
    blocks = DOCXParser().parse(path)
    list_blocks = [b for b in blocks if b.block_type_hint == "list_item"]
    assert len(list_blocks) == 3
